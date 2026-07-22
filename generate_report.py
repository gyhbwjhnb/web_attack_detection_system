"""
生成结题报告 .docx 文件
输出: docs/信息安全科技创新项目结题报告.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ==================== 全局样式统一 ====================

BODY_FONT = '宋体'
HEADING_FONT = '黑体'
BODY_SIZE = Pt(12)

# --- Normal 样式 ---
normal_style = doc.styles['Normal']
normal_style.font.name = BODY_FONT
normal_style.font.size = BODY_SIZE
normal_style.font.bold = False
normal_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
normal_style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', BODY_FONT)

# --- Heading 样式 ---
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = HEADING_FONT
    hs.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    hs.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', HEADING_FONT)
    if level == 1:
        hs.font.size = Pt(18)
    elif level == 2:
        hs.font.size = Pt(15)
    else:
        hs.font.size = Pt(13)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(text, bold=False, align=None, font_size=12):
    p = doc.add_paragraph()
    # 清除默认空段落
    p.clear()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # 设置东亚字体回退
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', BODY_FONT)
    if align:
        p.alignment = align
    return p


def add_screenshot_placeholder(description):
    """添加截图占位标记"""
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(f'【截图位置：{description}】')
    run.font.name = BODY_FONT
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def create_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(11)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(11)
    return table


# ==================== 封面 ====================
add_para("", font_size=12)
add_para("", font_size=12)
add_para("信息安全科技创新", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=26)
add_para("项目结题报告", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=26)
add_para("", font_size=12)
add_para("", font_size=12)

cover_items = [
    "项目名称：    网络攻击检测系统 (NIDS)",
    "项目组成员：   __________(姓名)  __________(班级)  __________(学号)",
    "               __________(姓名)  __________(班级)  __________(学号)",
    "               __________(姓名)  __________(班级)  __________(学号)",
    "院　　系：    __________",
    "学　　期：    2025-2026 学年第二学期",
    "报告完成日期：  2026年7月____日",
]
for item in cover_items:
    add_para(item, font_size=14)

doc.add_page_break()

# ==================== 项目摘要 ====================
add_heading("项目摘要", level=1)

add_para("""本项目设计并实现了一款基于特征匹配与异常行为分析双引擎的网络攻击检测系统（Network Intrusion Detection System, NIDS）。系统采用 Python 语言开发，图形界面基于 tkinter 内置库，无需额外第三方 GUI 依赖，可运行于 Windows/Linux/macOS 等平台。

项目围绕"抓包 → 检测 → 告警 → 展示"的核心数据流，采用四模块松耦合架构：模块一负责数据包捕获与协议解析，支持在线网卡抓包、离线 PCAP 回放和内存模拟注入三种数据源；模块二实现基于 Aho-Corasick 多模式匹配算法的特征检测引擎，内置 20 余类、30+ 条攻击特征规则，覆盖 SQL 注入、XSS、命令注入、WebShell、反弹 Shell 等常见 Web 攻击；模块三实现基于 EMA 自适应基线模型的异常行为检测引擎，搭载 10 个插件化检测器，可检测端口扫描、暴力破解、DDoS、异常外联、横向扩散、行为突变等异常流量；模块四提供完整的 tkinter 图形用户界面，支持实时告警展示、流量监控、多维筛选、统计面板、HTML/CSV/TXT 报告导出等功能。

在 v2.0 版本中，我们进一步引入三项智能增强特性：置信度动态评分（为 7 种攻击类型计算 0.0-1.0 置信分）、24 小时异常时段检测（基于 EMA 时段基线的凌晨活动感知）、自学习误报抑制（用户反馈驱动的跨会话持久化规则学习），使告警从"有无"升级为"可信度+时间感知+持续学习"的智能维度。

项目已完成全部设计、编码和测试工作，总代码量约 9,800 行 Python，包含 38 个源文件、7 个测试脚本，功能完整、运行稳定，已通过模拟攻击场景的全面验证。""")

doc.add_page_break()

# ==================== 目录 ====================
add_heading("目　录", level=1)
toc_items = [
    "项目摘要",
    "一、需求分析",
    "　　1.1 项目需要解决的安全问题",
    "　　1.2 项目功能目标",
    "二、总体设计",
    "　　2.1 总体架构图",
    "　　2.2 模块划分与介绍",
    "三、详细设计",
    "　　3.1 模块一：数据包捕获与预处理",
    "　　3.2 模块二：特征匹配检测引擎",
    "　　3.3 模块三：异常行为检测引擎",
    "　　3.4 模块四：图形用户界面",
    "四、系统实现与测试",
    "　　4.1 实现环境",
    "　　4.2 测试方法",
    "　　4.3 测试流程与结果",
    "五、项目总结",
    "　　5.1 实施概况",
    "　　5.2 不足与展望",
    "　　5.3 体会与建议",
    "附：项目组成员贡献表",
]
for item in toc_items:
    add_para(item, font_size=12)

doc.add_page_break()

# ==================== 一、需求分析 ====================
add_heading("一、需求分析", level=1)

add_heading("1.1 项目需要解决的安全问题", level=2)
add_para("""随着互联网的飞速发展，网络攻击手段日益多样化和自动化。传统防火墙仅能在网络层和传输层进行访问控制，无法检测应用层攻击（如 SQL 注入、XSS），也无法识别基于行为异常的未知攻击（如 APT 攻击的横向扩散阶段、C2 命令控制信道等）。

具体而言，本项目需要解决以下安全检测需求：""")
add_para("""(1) Web 应用攻击检测：能够识别 HTTP 流量中的 SQL 注入、跨站脚本(XSS)、命令注入、路径遍历、文件包含、WebShell 上传等应用层攻击。""")
add_para("""(2) 网络层异常检测：能够发现端口扫描、暴力破解、SYN Flood DDoS、异常外联、内网横向扩散等网络层攻击行为。""")
add_para("""(3) 未知行为感知：不仅依赖已知特征签名，还能基于历史行为基线发现偏离正常模式的异常流量。""")
add_para("""(4) 告警智能化：减少误报噪音，对告警进行可信度评估，并根据时段上下文动态调整告警等级。""")
add_para("""(5) 可视化与交互：提供直观的图形界面，支持实时监控、多维筛选、报告导出和规则自定义。""")

add_heading("1.2 项目功能目标", level=2)
add_para("""本项目的核心功能目标包括：""")
add_para("""(1) 多源数据输入：支持在线网卡抓包（需 Npcap）、离线 PCAP 文件回放、内存模拟数据注入三种模式。""")
add_para("""(2) 双引擎并行检测：特征匹配引擎（基于 Aho-Corasick 多模式匹配，覆盖 20+ 类 Web 攻击）和异常行为引擎（基于 EMA 自适应基线 + 10 个插件化检测器）同时运行。""")
add_para("""(3) 插件化架构：检测器通过 IDetector 接口实现，新增检测器仅需 3 步（新建文件 → 继承接口 → 注册），无需修改引擎核心代码。""")
add_para("""(4) 规则热加载：在 data/rules/ 目录下拖入 JSON 规则文件，系统 3 秒内自动检测并重新加载，无需重启。""")
add_para("""(5) 智能告警增强：置信度动态评分（7 种攻击类型 0.0-1.0 评分）、24h 时段基线（凌晨异常自动提权）、自学习误报抑制（用户反馈 → 持久化规则 → 自动降权）。""")
add_para("""(6) 完整 GUI 功能：告警/流量双表格实时展示，按严重度/协议着色，三维 IP/协议/端口组合筛选，统计面板（攻击类型分布/严重度分布/攻击链），HTML/CSV/TXT 报告导出，白名单管理，6 标签页帮助系统。""")

# ==================== 二、总体设计 ====================
add_heading("二、总体设计", level=1)

add_heading("2.1 总体架构图", level=2)
add_para("""系统的总体架构遵循"四模块松耦合 + MessageBus 消息总线"的设计模式。各模块之间不直接依赖，仅通过全局 MessageBus（发布-订阅模式）进行数据交换，从而实现高度的模块独立性和可扩展性。""")

add_screenshot_placeholder("总体架构图 — 请截取 docs/项目报告.md 中第 2.1 节的 ASCII 架构图，或运行程序截图主界面")

add_para("""系统由四个核心模块组成：""")
add_para("""模块一（数据捕获与预处理）：负责从网卡/PCAP/模拟源获取原始数据包，经协议解析器转换为统一的 TrafficRecord 数据结构，通过 MessageBus 发布 traffic_record 事件。""")
add_para("""模块二（特征匹配检测引擎）：订阅 traffic_record 事件，基于 Aho-Corasick 多模式匹配算法，使用 data/signatures.json 中的 30+ 条攻击特征规则进行检测，命中后发布 signature_alert 事件。支持规则热加载。""")
add_para("""模块三（异常行为检测引擎）：订阅 traffic_record 事件，通过 EMA 自适应基线模型学习主机行为模式，10 个插件化检测器按优先级遍历，检测异常行为后发布 anomaly_alert 事件。后处理阶段执行置信度评分和误报抑制。""")
add_para("""模块四（图形用户界面）：订阅 signature_alert 和 anomaly_alert 事件，在 tkinter 表格中实时展示告警和流量数据，提供筛选、统计、导出、白名单、帮助等完整功能。""")

add_heading("2.2 模块划分与介绍", level=2)

add_para("【模块一：数据包捕获与预处理】", bold=True)
add_para("功能：负责从三种数据源获取原始流量并解析为统一的 TrafficRecord 格式。")
add_para("输入：网卡原始数据包 / PCAP 文件 / 模拟生成的 scapy 包。")
add_para("输出：TrafficRecord（包含 src/dst IP+端口、协议、flags、payload、HTTP 字段、DNS 字段、TLS SNI 等）。")

add_para("【模块二：特征匹配检测引擎】", bold=True)
add_para("功能：基于已知攻击特征签名的模式匹配检测。")
add_para("输入：TrafficRecord + data/signatures.json 规则库。")
add_para("输出：Alert（攻击类型、严重度、源/目标 IP、匹配载荷片段、处置建议）。")

add_para("【模块三：异常行为检测引擎】", bold=True)
add_para("功能：基于 EMA 行为基线的异常流量检测。")
add_para("输入：TrafficRecord + 历史基线数据（hosts/baselines）。")
add_para("输出：Alert（含置信度评分）+ AttackChain（攻击链事件）。")

add_para("【模块四：图形用户界面】", bold=True)
add_para("功能：提供实时监控、筛选、统计、导出等交互功能。")
add_para("输入：signature_alert / anomaly_alert / statistics 事件。")
add_para("输出：tkinter GUI 界面、HTML/CSV/TXT 报告文件。")

add_para("", font_size=12)
add_para("模块间信息交互通过 MessageBus 的 6 个标准事件完成：traffic_record（模块1→模块2/3/4）、signature_alert（模块2→模块4）、anomaly_alert（模块3→模块4）、attack_chain（模块3→模块4）、statistics（模块3→模块4）、config_change（模块4→模块2/3）。")

add_screenshot_placeholder("MessageBus 事件流图 — 请截取程序运行时的数据流动示意，或使用项目报告中的流程图")

doc.add_page_break()

# ==================== 三、详细设计 ====================
add_heading("三、详细设计", level=1)

# --- 3.1 模块一 ---
add_heading("3.1 模块一：数据包捕获与预处理", level=2)
add_para("""模块一作为系统的唯一数据入口，核心类为 CaptureEngine（capture.py），实现 ICaptureEngine 接口。支持三种工作模式：在线抓包模式使用 scapy.sniff() + BPF 过滤实时捕获网卡数据；离线模式使用 scapy.rdpcap() 读取 PCAP 文件；模拟模式直接构造 TrafficRecord 注入 MessageBus，无需 Npcap 和管理员权限。""")

add_para("核心数据结构 TrafficRecord 包含以下主要字段：", bold=True)
create_table(
    ["字段名", "类型", "说明"],
    [
        ["id", "str", "唯一标识（时间戳+随机数）"],
        ["timestamp", "float", "数据包捕获时间戳"],
        ["src / dst", "IPEndpoint", "源/目标端点（IP + port）"],
        ["protocol", "ProtocolType", "TCP / UDP / ICMP / HTTP / DNS / TLS / ARP"],
        ["flags", "int", "TCP Flags（SYN=0x02, PSH+ACK=0x18 等）"],
        ["payload", "str", "应用层载荷文本"],
        ["http_method/uri/host/ua", "str", "HTTP 协议字段"],
        ["dns_query/qtype", "str/int", "DNS 查询字段"],
        ["tls_sni", "str", "TLS SNI 字段"],
        ["app_protocol", "str", "自动识别的应用层协议"],
    ],
)
add_para("", font_size=6)

add_para("关键函数设计：", bold=True)
create_table(
    ["函数名", "输入", "输出", "功能"],
    [
        ["parse_packet()", "scapy Packet", "TrafficRecord", "逐层解析以太帧→IP→TCP/UDP→HTTP/DNS/TLS/ARP"],
        ["CaptureEngine.start()", "mode, source", "None", "启动抓包线程，根据模式选择数据源"],
        ["CaptureEngine.stop()", "None", "None", "停止抓包，清理资源"],
        ["create_fake_http_record()", "method, host, uri", "TrafficRecord", "生成模拟 HTTP 流量（用于测试）"],
    ],
)
add_para("", font_size=6)
add_screenshot_placeholder("模块一代码文件结构截图（module1_capture/ 目录）")

# --- 3.2 模块二 ---
add_heading("3.2 模块二：特征匹配检测引擎", level=2)
add_para("""模块二实现基于已知攻击特征的签名匹配检测，核心类为 SignatureEngine（signature_engine.py），配合 AhoCorasickMatcher（matcher.py）多模式匹配器。""")

add_para("""Aho-Corasick 多模式匹配算法是本模块的核心：通过构建 Trie 树 + BFS 构建 Failure Link（失败指针），实现在 O(n) 时间内同时匹配所有规则模式（n 为待检测文本长度），相比逐条正则匹配效率提升数十倍。匹配器还实现了 URL 解码归一化对抗绕过（如 %20→空格、%27→单引号），防止攻击者通过编码绕过检测。""")

add_para("检测流程：", bold=True)
add_para("""(1) 提取检测文本：优先取 HTTP 完整字段（method + host + uri + user_agent），无 HTTP 字段则取 payload。""")
add_para("""(2) Aho-Corasick 一次扫描：在检测文本中匹配所有已注册规则的模式串。""")
add_para("""(3) 三级过滤：协议过滤（TCP/UDP 匹配）、端口过滤（dst_port/src_port 匹配）、正则二次验证（如规则含 pattern 字段）。""")
add_para("""(4) 暴力破解特殊处理：统计同一 (src_ip, dst_ip, dst_port) 在窗口内的 SYN 包数量，超过阈值生成暴力破解告警。""")
add_para("""(5) 告警去重：同一 (src_ip, dst_ip, attack_type) 在窗口期内不重复告警。""")

add_para("", font_size=6)
add_para("规则热加载机制：", bold=True)
add_para("""start_hot_reload("data/rules/") 启动后台线程，每 3 秒轮询目录下 .json 文件的修改时间，检测到变化后自动调用 load_rules() 重新加载全部规则，控制台输出"检测到规则变更，已重新加载 N 条规则"确认信息。相对路径自动通过 _project_root（signature_engine.py 所在目录的上一级 = 项目根目录）解析为绝对路径。""")

add_screenshot_placeholder("特征规则 JSON 示例截图（data/signatures.json 部分内容）")

# --- 3.3 模块三 ---
add_heading("3.3 模块三：异常行为检测引擎", level=2)
add_para("""模块三实现基于 EMA 自适应基线模型的异常行为检测，核心类为 AnomalyEngine（anomaly_engine.py），检测逻辑由 10 个独立插件的 IDetector 子类实现。""")

add_para("EMA 基线模型：", bold=True)
add_para("""每台主机的连接数、对端 IP 数、带宽使用量、SYN 比例等指标采用指数加权移动平均（EMA）更新：新基线 = α × 当前值 + (1-α) × 旧基线，α = 2/(N+1)，N ≈ 60（约 2 分钟窗口）。EMA 使基线既不受瞬时波动影响，又能逐步适应长期行为变化（如工作时间 vs 非工作时间的流量差异）。""")

add_para("", font_size=6)
add_para("10 个检测器一览：", bold=True)
create_table(
    ["优先级", "检测器名称", "category", "检测逻辑"],
    [
        ["10", "PortScanDetector", "scan", "单 IP 窗口内访问 N 个不同端口 > 阈值"],
        ["15", "SynFloodDetector", "ddos", "窗口内 SYN 包总数 > 阈值"],
        ["18", "BandwidthAnomalyDetector", "traffic", "带宽 > 基线 EMA × 倍数"],
        ["20", "BruteForceDetector", "brute", "敏感端口(22/3389/3306)连接密度 > 阈值"],
        ["25", "AbnormalOutboundDetector", "exfil", "连接外网陌生 IP + 恶意/非标准端口"],
        ["30", "LateralMovementDetector", "lateral", "内网 IP 连接多个内网 IP + 跨子网"],
        ["40", "BehavioralDeviationDetector", "behavior", "4维度(连接数/对端/端口/SYN)偏离基线 > 倍率"],
        ["50", "TimeProfileDetector", "time", "24h EMA 基线，凌晨连接 > 均值+2σ → 置信度提权"],
        ["80", "SuppressionLearner", "filter", "用户标记误报 → 学习签名 → 后续同模式自动降权"],
        ["85", "ConfidenceScorer", "scoring", "7种攻击类型动态计算 0.0-1.0 置信度评分"],
    ],
)
add_para("", font_size=6)

add_para("IDetector 插件接口：", bold=True)
add_para("""所有检测器必须实现 IDetector（common/detector.py）抽象基类。name 字段为唯一标识（下划线命名），priority 字段决定执行顺序（10-50 主检测阶段，80-90 后处理阶段）。必须实现 process(record, now) → List[Alert] 方法，可选实现 post_process_batch(alerts) → List[Alert] 后处理方法。""")
add_para("""AnomalyEngine 在启动时遍历 _ALL_DETECTORS 列表实例化所有检测器，调用 set_context() 注入共享上下文（hosts/baselines/config/lock/count_recent/make_alert），按 priority 升序排序。处理每条流量时，依次调用每个检测器的 process() 方法，然后遍历调用 post_process_batch() 进行后处理。""")

add_screenshot_placeholder("检测器代码示例截图（任意一个 detector.py 文件）")
add_screenshot_placeholder("IDetector 接口代码截图（common/detector.py）")

# --- 3.4 模块四 ---
add_heading("3.4 模块四：图形用户界面", level=2)
add_para("""模块四为基于 tkinter 的完整图形用户界面，核心类为 MainWindow（main_window.py），单文件约 1,750 行代码。""")

add_para("主要功能类（均为 MainWindow 内部方法）：", bold=True)
create_table(
    ["方法名", "功能", "说明"],
    [
        ["_add_alert_to_table()", "告警入表", "按严重度5级着色（蓝→红），新数据自动追加到行末"],
        ["_add_traffic_to_table()", "流量入表", "按协议4色标记（TCP/UDP/HTTP/DNS），支持 bbox 检测自动滚动"],
        ["_apply_filter()", "三维筛选", "IP/协议/端口组合过滤，带历史值下拉建议"],
        ["_update_stats_panel()", "统计更新", "告警总数、攻击类型分布、严重度分布（进度条）、攻击链"],
        ["_export_html_report()", "HTML报告", "含概览卡片 + 分布图 + 明细表的完整分析报告"],
        ["_export_csv()", "CSV导出", "告警数据导出为 CSV 文件（Excel 可打开）"],
        ["_open_help_dialog()", "帮助系统", "6标签页：快速开始/界面/筛选/白名单/导出/自定义规则"],
        ["_on_import_rules()", "规则导入", "文件对话框选择 JSON → 复制到 data/rules/ → 触发热加载"],
    ],
)
add_para("", font_size=6)
add_screenshot_placeholder("GUI 主界面完整截图（运行 python main.py 后截图）")
add_screenshot_placeholder("帮助弹窗—自定义规则标签页截图")

doc.add_page_break()

# ==================== 四、系统实现与测试 ====================
add_heading("四、系统实现与测试", level=1)

add_heading("4.1 实现环境", level=2)
create_table(
    ["项目", "说明"],
    [
        ["编程语言", "Python 3.10"],
        ["GUI 框架", "tkinter（Python 内置，无需额外安装）"],
        ["网络抓包", "scapy（唯一第三方依赖，pip install scapy）"],
        ["文档生成", "python-pptx（PPT 生成）、python-docx（报告生成）"],
        ["操作系统", "Windows 11（兼容 Linux/macOS，仅抓包需 Npcap）"],
        ["开发工具", "VS Code / Trae IDE"],
        ["版本管理", "Git + GitHub"],
        ["打包工具", "PyInstaller（生成独立 .exe）"],
    ],
)

add_heading("4.2 测试方法与流程", level=2)
add_para("""本项目采用分层测试策略，包含单元测试、集成测试、实机演示和智能特性验证四个层面：""")
add_para("""(1) 单元测试（test_module1/2/3.py）：针对每个模块的核心功能编写独立测试用例。模块1 测试验证 Scapy 包构造 → TrafficRecord 解析的正确性；模块2 测试验证 Aho-Corasick 多模式匹配算法的正确性和 SignatureEngine 的规则匹配准确性；模块3 测试验证 10 个检测器插件的异常检测逻辑。""")
add_para("""(2) 集成测试（quick_test.py）：构造模拟攻击流量注入 MessageBus，验证从流量输入 → 双引擎检测 → 告警输出的完整数据流通路。""")
add_para("""(3) 实机演示（live_demo.py）：在独立线程中持续生成模拟攻击流量（包含端口扫描、SQL注入、XSS、暴力破解等多种攻击），通过 GUI 实时观察检测效果，适合汇报展示。""")
add_para("""(4) 智能特性验证（test_intelligence.py）：针对 v2.0 新增的 S6 置信度评分、S7 时段检测、S8 自学习误报抑制三项智能增强功能，模拟 6 个攻击场景逐一验证。""")

add_para("", font_size=6)
add_para("测试流程：", bold=True)
add_para("""① 先注入 30 条正常 HTTP/HTTPS 流量（google.com），建立行为基线。""")
add_para("""② 等待 5 秒学习期结束后，依次注入 6 个攻击场景的流量：端口扫描（80 端口/10s）、SSH 暴力破解（50次/30s）、SYN Flood（500包/3s）、夜间 C2 外联（凌晨3点/20次）、异常外联（15个外网IP）、横向扩散（12个内网主机）。""")
add_para("""③ 在 GUI 中观察告警列表，验证：每条告警的 confidence 列数值是否正确变化（不同攻击类型分值不同）；夜间 C2 告警描述是否包含"[夜间异常时段]"标记；右键标记某告警为"误报"后，后续同类型告警的 confidence 是否自动降低。""")

add_screenshot_placeholder("测试运行截图 — 终端输出（python test_intelligence.py 的控制台）")
add_screenshot_placeholder("测试结果截图 — GUI 告警列表（展示 confidence 列和告警描述）")

add_heading("4.3 测试结果与结论", level=2)
add_para("""经过全面测试，系统运行稳定，各项功能均达到设计目标：""")
add_para("""(1) 特征匹配引擎：能正确检测 SQL 注入、XSS、命令注入、路径遍历、WebShell 等 20+ 类已知攻击；Aho-Corasick 多模式匹配 + URL 解码对抗可有效应对编码绕过攻击。""")
add_para("""(2) 异常检测引擎：10 个插件化检测器均可正常触发告警，EMA 基线模型能区分正常流量波动与真正的异常行为；置信度评分使不同严重程度的告警具有不同的可信度分值（端口扫描 80 端口可达 0.92，仅扫 3 端口约 0.33）。""")
add_para("""(3) 智能增强：时段检测能正确识别凌晨非工作时间的异常活动并自动提升置信度；误报抑制在用户标记 2 次后自动生效，后续同模式告警置信度降低 0.35。""")
add_para("""(4) GUI 功能：告警/流量双表格实时更新正常，筛选栏正确过滤数据，HTML/CSV/TXT 报告导出格式正确，白名单即时生效，6 标签页帮助系统内容完整。""")

doc.add_page_break()

# ==================== 五、项目总结 ====================
add_heading("五、项目总结", level=1)

add_heading("5.1 项目实施概况", level=2)
add_para("""本项目从需求分析到编码实现再到测试完善，历时一个学期。项目采用迭代式开发，共经历两个大版本：v1.0 完成了四模块基础架构、三种数据源、双引擎检测和基本 GUI 功能；v2.0 重构为插件化检测器架构，新增 7 个独立检测器文件 + IDetector 接口，引入规则热加载、置信度评分、时段检测和自学习误报抑制四项增强。""")
add_para("""项目总代码量约 9,800 行（38 个 Python 文件 + 7 个测试脚本），文档约 3,000 行（README + 项目报告 + 设计文档 + 项目介绍 + PPT），已推送到 GitHub 开源仓库。独立 exe 安装包约 15MB，可在无 Python 环境的 Windows 系统上直接运行。""")

add_heading("5.2 不足与展望", level=2)
add_para("""当前版本存在的不足：""")
add_para("""(1) 性能瓶颈：Python 的单线程 GIL 限制导致高流量场景下（>100Mbps）可能出现丢包，如需生产级部署建议重写核心检测模块为 C++。""")
add_para("""(2) 检测器覆盖：目前 10 个异常检测器覆盖了主要攻击面，但缺少 C2 Beacon 心跳检测、DNS 隧道深度检测、载荷熵值分析等高级检测能力。""")
add_para("""(3) GUI 技术栈：tkinter 功能有限，界面美观度不如 Web 前端框架（如 React + ECharts）。""")
add_para("""(4) 无机器学习：当前所有检测基于规则和统计，未引入监督/无监督学习算法（如 Isolation Forest 异常检测）。""")

add_para("", font_size=6)
add_para("未来可扩展方向：", bold=True)
add_para("""(1) 新增 5 个检测器：C2 Beacon Detector（周期性心跳）、Payload Entropy Detector（载荷熵值）、DNS Tunnel Detector（DNS 隧道）、Alert Correlator（告警关联聚合）、Protocol Masquerade Detector（协议伪装）。""")
add_para("""(2) 告警输出插件化：实现 IAlertSink 接口，支持 Syslog/Email/Slack Webhook/Elasticsearch 等多通道告警推送。""")
add_para("""(3) 机器学习引擎（模块5）：引入 Isolation Forest / Autoencoder 无监督异常检测，弥补规则类检测对零日攻击的盲区。""")
add_para("""(4) Web Dashboard：使用 Flask + ECharts 构建浏览器端监控界面，替代 tkinter 实现远程多终端访问。""")

add_heading("5.3 体会与建议", level=2)
add_para("""通过本项目，我们对网络入侵检测系统的设计和实现有了全面深入的理解。主要体会包括：""")
add_para("""(1) 架构是核心：良好的松耦合架构（MessageBus + 插件化）使后期功能扩展只需新增文件而无需修改核心代码，大大降低了维护成本。""")
add_para("""(2) 双引擎必要性：纯特征匹配无法发现未知攻击，纯异常检测又会产生大量误报——两者结合实现了"已知攻击精准拦"和"异常行为不放过"的互补。""")
add_para("""(3) 告警不止于"有/无"：从简单的"检测到攻击"升级为"置信度 0.92 + 凌晨异常 + 已学习不再误报"的智能告警，显著提升了安全运营效率。""")
add_para("""(4) 测试驱动开发：每条检测器都应该有对应的测试场景和验证流程，test_intelligence.py 的六场景测试编写成本极低但收益巨大。""")

add_para("", font_size=6)
add_para("建议与意见：", bold=True)
add_para("""(1) 建议课程增加自动化测试与 CI/CD 的课时，帮助同学们建立测试驱动的开发习惯。""")
add_para("""(2) 建议引入 Git 协作工作流的实训内容（如 Pull Request 流程、Code Review 实践）。""")
add_para("""(3) 建议课程期末设置项目答辩环节，通过 PPT 汇报促进同学们总结和表达能力。""")

doc.add_page_break()

# ==================== 附录：贡献表 ====================
add_heading("附：项目组成员贡献表", level=1)
add_para("（由项目组长填写，如不填写，各成员平分工作量）", font_size=11)

create_table(
    ["成员姓名", "是否项目组长", "具体承担任务", "组长评分（百分制）"],
    [
        ["__________", "是/否", "__________", "_____"],
        ["__________", "是/否", "__________", "_____"],
        ["__________", "是/否", "__________", "_____"],
        ["__________", "是/否", "__________", "_____"],
    ],
)

# ==================== 保存 ====================
output_path = "docs/信息安全科技创新项目结题报告.docx"
os.makedirs("docs", exist_ok=True)
doc.save(output_path)
print(f"结题报告已生成: {output_path}")
